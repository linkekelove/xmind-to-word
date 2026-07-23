"""
XMind 转 Word 工具
解析 .xmind 文件（ZIP 格式），按原有父子节点层级关系生成 Word 文档，
并应用中文排版习惯的字体、段落格式。

用法:
    python xmind_to_word.py <输入.xmind> <输出.docx>
"""

import sys
import json
import re
import zipfile
import tempfile
import shutil
import os

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

MAX_HEADING = 9

# 中文排版字号（磅）
SIZE_TITLE = 22   # 二号（封面标题）
SIZE_H1 = 22       # 二号
SIZE_H2 = 15       # 小三
SIZE_H3 = 14       # 四号
SIZE_H4 = 12       # 小四（与正文同号，靠黑体+加粗区分层级）
SIZE_BODY = 12     # 小四

FONT_HEADING = '黑体'
FONT_BODY = '宋体'

# 标题1~4 用黑体区分层级，标题5及以下与正文一样按正文样式处理
HEADING_STYLES = {'Title', 'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4'}
BODY_PARAGRAPH_SPACE_AFTER = 12  # 正文段后间距（磅），近似一行高度


def extract_xmind(xmind_path, extract_dir):
    """解压 .xmind 文件到指定目录"""
    with zipfile.ZipFile(xmind_path, 'r') as zf:
        zf.extractall(extract_dir)


# 全角空格(U+3000)、不间断空格(U+00A0)、零宽空格(U+200B)、
# 零宽连字/断字(U+200C/200D)、BOM(U+FEFF) 等不可见排版污染字符
_INVISIBLE_CHARS = re.compile(
    "[\u3000\u00a0\u200b\u200c\u200d\ufeff]"
)
_CONTROL_CHARS = re.compile("[\x00-\x08\x0b\x0c\x0e-\x1f]")


def clean_text(text):
    """清除全角空格 / 不间断空格 / 零宽字符 / 控制字符等排版污染"""
    text = _INVISIBLE_CHARS.sub(' ', text)
    text = _CONTROL_CHARS.sub('', text)
    return re.sub(r' {2,}', ' ', text).strip()


def parse_topic(topic, level=0):
    """递归解析思维导图节点，保留父子层级关系"""
    result = []

    title = clean_text(topic.get('title', ''))
    if title:
        result.append({'level': level, 'text': title})

    image = topic.get('image')
    if image:
        src = image.get('src', '')
        if src.startswith('xap:'):
            src = src[4:]
        result.append({'level': level, 'image': src})

    children = topic.get('children', {})
    for child in children.get('attached', []):
        result.extend(parse_topic(child, level + 1))

    return result


def set_run_font(run, font_name, size_pt, bold=False):
    """设置字体（中英文字体、字号、颜色）"""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = False
    run.font.color.rgb = RGBColor(0, 0, 0)

    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def set_first_line_indent_chars(paragraph, chars=2):
    """设置首行缩进（按字符数，中文排版惯用单位）"""
    pPr = paragraph._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = pPr.makeelement(qn('w:ind'), {})
        pPr.append(ind)
    ind.set(qn('w:firstLineChars'), str(chars * 100))
    if ind.get(qn('w:firstLine')) is not None:
        del ind.attrib[qn('w:firstLine')]


_CN_DIGITS = '零一二三四五六七八九'
_CN_UNITS = ['', '十', '百', '千']


def chinese_numeral(n):
    """将 1~9999 的整数转换为中文数字（用于二级标题编号，如 一、二、三）"""
    if n <= 0:
        return _CN_DIGITS[0]
    if 10 <= n < 20:
        return '十' + (_CN_DIGITS[n % 10] if n % 10 else '')

    s = ''
    zero_flag = False
    unit_pos = 0
    num = n
    while num > 0:
        digit = num % 10
        if digit == 0:
            if s and not zero_flag:
                s = '零' + s
                zero_flag = True
        else:
            s = _CN_DIGITS[digit] + _CN_UNITS[unit_pos] + s
            zero_flag = False
        unit_pos += 1
        num //= 10
    if s.startswith('一十'):
        s = s[1:]
    return s


def apply_numbering(items):
    """
    按父子层级关系为标题添加序号：
    - 一级标题(H1)：文档中只有唯一一个根节点，不需要区分编号，跳过
    - 二级标题(H2)：中文数字，如 一、二、三
    - 三级标题(H3)：两级阿拉伯数字，如 1.1、1.2（首位对应所属二级标题的序号）
    - 四级标题(H4)：三级阿拉伯数字，如 1.1.1
    - 五级及以下：不编号，保留原有项目符号缩进
    """
    counters = {2: 0, 3: 0, 4: 0}

    for item in items:
        if 'text' not in item:
            continue  # 图片节点不参与编号

        heading_level = item['level'] + 1

        if heading_level == 1:
            continue  # 唯一的一级标题，不编号
        elif heading_level == 2:
            counters[2] += 1
            counters[3] = 0
            counters[4] = 0
            label = f'{chinese_numeral(counters[2])}、'
        elif heading_level == 3:
            counters[3] += 1
            counters[4] = 0
            label = f'{counters[2]}.{counters[3]}  '
        elif heading_level == 4:
            counters[4] += 1
            label = f'{counters[2]}.{counters[3]}.{counters[4]}  '
        else:
            continue  # 五级及以下不编号

        item['text'] = label + item['text']


def reset_char_spacing(run):
    """将字符间距显式重置为标准（0），避免继承样式或外部来源写入的加宽值"""
    rpr = run._element.get_or_add_rPr()
    spacing = rpr.find(qn('w:spacing'))
    if spacing is None:
        spacing = rpr.makeelement(qn('w:spacing'), {})
        rpr.append(spacing)
    spacing.set(qn('w:val'), '0')


def create_word_doc(content_list, output_path, resource_dir, doc_title='文档'):
    """根据解析结果生成 Word 文档，并应用中文排版格式"""
    doc = Document()

    title_p = doc.add_heading(doc_title, 0)
    title_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for item in content_list:
        level = item['level']

        if 'image' in item:
            img_path = os.path.join(resource_dir, item['image'])
            try:
                doc.add_picture(img_path, width=Inches(5))
            except Exception:
                doc.add_paragraph(f'[图片加载失败: {item["image"]}]')
            continue

        text = item['text']
        heading_level = level + 1

        if heading_level <= MAX_HEADING:
            p = doc.add_heading(text, level=heading_level)
        else:
            indent = '    ' * (heading_level - MAX_HEADING - 1)
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(indent + text)

        apply_style(p)

    doc.save(output_path)


def apply_style(paragraph):
    """对单个段落应用字体、颜色、对齐、间距、缩进等中文排版格式"""
    style_name = paragraph.style.name

    if style_name == 'Heading 1':
        font, size, bold = FONT_HEADING, SIZE_H1, True
    elif style_name == 'Heading 2':
        font, size, bold = FONT_HEADING, SIZE_H2, True
    elif style_name == 'Heading 3':
        font, size, bold = FONT_HEADING, SIZE_H3, True
    elif style_name == 'Heading 4':
        font, size, bold = FONT_HEADING, SIZE_H4, True
    else:
        font, size, bold = FONT_BODY, SIZE_BODY, False

    for run in paragraph.runs:
        set_run_font(run, font, size, bold=bold)
        reset_char_spacing(run)

    if style_name not in HEADING_STYLES:
        pf = paragraph.paragraph_format
        # 短语类节点内容两端对齐会被强行拉伸字间距撑满整行，改为左对齐
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.space_before = Pt(0)
        pf.space_after = Pt(BODY_PARAGRAPH_SPACE_AFTER)
        pf.line_spacing = 1.0
        pf.left_indent = 0
        set_first_line_indent_chars(paragraph, 2)


def convert(xmind_path, output_path):
    """主转换流程：解压 -> 解析 -> 生成 Word"""
    tmp_dir = tempfile.mkdtemp(prefix='xmind_')
    try:
        extract_xmind(xmind_path, tmp_dir)

        with open(os.path.join(tmp_dir, 'content.json'), 'r', encoding='utf-8') as f:
            data = json.load(f)

        all_items = []
        for sheet in data:
            root_topic = sheet.get('rootTopic', {})
            all_items.extend(parse_topic(root_topic, 0))

        apply_numbering(all_items)

        doc_title = os.path.splitext(os.path.basename(xmind_path))[0]
        create_word_doc(all_items, output_path, tmp_dir, doc_title=doc_title)

        print(f'共解析 {len(all_items)} 个节点')
        print(f'Word 文档已保存到: {output_path}')
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


def main():
    if len(sys.argv) != 3:
        print('用法: python xmind_to_word.py <输入.xmind> <输出.docx>')
        sys.exit(1)

    convert(sys.argv[1], sys.argv[2])


if __name__ == '__main__':
    main()
